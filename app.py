import streamlit as st
import google.generativeai as genai
import json
import os
import time
from docx import Document


genai.configure(
    api_key=st.secrets["GEMINI_API_KEY"]
)


st.set_page_config(
    page_title="深度文学创作引擎",
    layout="wide",
    initial_sidebar_state="expanded"
)


SAVE_FILE="novel_save.json"



st.markdown("""
<style>

.stChatMessage {
font-size:13px!important;
line-height:1.8!important;
}

p,li,span,div[data-testid="stMarkdownContainer"]{
font-size:13px!important;
line-height:1.8!important;
}

.stTextArea textarea{
font-size:14px!important;
min-height:200px!important;
}

.stButton>button{
width:100%;
}

</style>
""",
unsafe_allow_html=True)



# =====================
# 保存系统
# =====================


def load_data():

    if os.path.exists(SAVE_FILE):

        try:

            with open(
                SAVE_FILE,
                "r",
                encoding="utf-8"
            ) as f:

                return json.load(f)

        except:

            return {}

    return {}




def sync_data():

    data={

        "chat_history":
        st.session_state.chat_history,

        "bible_info":
        st.session_state.bible_info,

        "pre_text":
        st.session_state.pre_text,

        "style_info":
        st.session_state.style_info,

        "ml_info":
        st.session_state.ml_info,

        "fl_info":
        st.session_state.fl_info,

        "bg_info":
        st.session_state.bg_info

    }


    with open(
        SAVE_FILE,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            data,
            f,
            ensure_ascii=False,
            indent=4
        )





def export_docx():

    doc=Document()

    doc.add_heading(
        "小说正文",
        level=1
    )


    for msg in st.session_state.chat_history:

        if msg["role"]=="model":

            doc.add_paragraph(
                msg["text"]
            )


    path="novel_export.docx"

    doc.save(path)

    return path





# =====================
# API保护
# =====================


def safe_send(chat,prompt):

    for i in range(5):

        try:

            return chat.send_message(prompt)


        except Exception as e:

            if "429" in str(e) or "ResourceExhausted" in str(e):

                st.warning(
                    f"请求过多，等待 {i+1} 次重试"
                )

                time.sleep(
                    (i+1)*8
                )

            else:

                raise e


    return None






# =====================
# 初始化
# =====================


data=load_data()


if "chat_history" not in st.session_state:


    st.session_state.chat_history=data.get(
        "chat_history",
        []
    )


    st.session_state.bible_info=data.get(
        "bible_info",
        ""
    )


    st.session_state.pre_text=data.get(
        "pre_text",
        ""
    )


    st.session_state.style_info=data.get(
        "style_info",
        ""
    )


    st.session_state.ml_info=data.get(
        "ml_info",
        ""
    )


    st.session_state.fl_info=data.get(
        "fl_info",
        ""
    )


    st.session_state.bg_info=data.get(
        "bg_info",
        ""
    )




st.title(
    "✍️ 深度文学创作引擎"
)





# =====================
# 侧边栏
# =====================


with st.sidebar:


    st.header(
        "🛡️ 数据中心"
    )


    if st.button(
        "💾 保存当前"
    ):

        sync_data()

        st.success(
            "已保存"
        )



    export=json.dumps(
        st.session_state.chat_history,
        ensure_ascii=False,
        indent=4
    )


    st.download_button(
        "📥 导出JSON备份",
        export,
        "novel_backup.json"
    )



    if st.button(
        "📄生成DOCX小说"
    ):

        file=export_docx()

        with open(
            file,
            "rb"
        ) as f:

            st.download_button(
                "下载DOCX",
                f,
                "小说.docx"
            )



    if st.button(
        "🗑️清空全部"
    ):

        st.session_state.chat_history=[]

        st.session_state.bible_info=""

        sync_data()

        st.rerun()

# =====================
# 设定区域
# =====================


with st.expander(
    "📚 核心档案与风格学习",
    expanded=True
):


    st.session_state.bible_info=st.text_area(
        "【核心档案库】剧情记忆:",
        st.session_state.bible_info,
        height=180
    )


    st.session_state.pre_text=st.text_area(
        "【参考例文】输入你的文风样本:",
        st.session_state.pre_text,
        height=220
    )


    st.session_state.style_info=st.text_area(
        "【文风档案】AI分析结果:",
        st.session_state.style_info,
        height=180
    )


    c1,c2=st.columns(2)


    st.session_state.ml_info=c1.text_input(
        "男主设定:",
        st.session_state.ml_info
    )


    st.session_state.fl_info=c2.text_input(
        "女主设定:",
        st.session_state.fl_info
    )


    st.session_state.bg_info=c1.text_input(
        "背景设定:",
        st.session_state.bg_info
    )



    if st.button(
        "💾 保存所有设置"
    ):

        sync_data()

        st.success(
            "设置已保存"
        )





# =====================
# 参考例文分析
# =====================


if st.button(
    "🧠 分析参考例文生成文风档案"
):

    try:

        editor=genai.GenerativeModel(
            "gemini-2.5-flash"
        )


        result=editor.generate_content(
f"""
你是一名专业小说编辑。

分析下面参考文本。

不要续写。

生成【文风档案】。


分析：

1.叙事视角
2.句式节奏
3.语言风格
4.描写密度
5.环境描写特点
6.人物动作特点
7.心理描写方式
8.对白习惯
9.情绪推进方式
10.特殊写作习惯


参考文本：

{st.session_state.pre_text[-4000:]}
"""
        )


        st.session_state.style_info=result.text

        sync_data()

        st.success(
            "文风档案生成完成"
        )


    except Exception as e:

        st.error(
            f"分析失败:{e}"
        )





# =====================
# 核心文学指令
# =====================


system_prompt=f"""

你是一名顶级商业小说作者。

你的任务：

只输出小说正文。

不要解释。

不要聊天。

不要回答用户。


【最高规则】

用户输入的是剧情导演指令。

不是正文。


禁止：

复制用户提示词。

复述用户要求。

总结用户要求。


禁止出现：

好的

以下是

根据你的要求

续写如下

我将为你


必须像小说书页一样，
直接进入场景。



【参考文风】

必须学习用户提供的例文：

句式

节奏

叙事视角

语言习惯

情绪表达

对白风格



【文学描写强制要求】


禁止流水账。


每个场景必须展开。



【环境描写】

必须加入：

天气

光线

空间

声音

气味

温度

触感


环境必须影响人物情绪。



【动作描写】

禁止：

他很生气。

她很难过。


必须通过：

动作

姿态

手指

呼吸

眼神

停顿

反应

表现。



【微表情】

加入：

眉眼变化

嘴角

目光停留

脸色变化

下意识动作



【心理描写】

加入：

犹豫

挣扎

克制

欲望

恐惧

判断

回忆



【对白】

对白不能只是传递信息。

必须体现：

关系

冲突

试探

隐藏目的

情绪。



【感官】

强化：

视觉

听觉

嗅觉

触觉



【节奏】

重要剧情慢写。

关键事件：

动作

情绪爆发

关系变化

不得一句带过。



男主：

{st.session_state.ml_info}


女主：

{st.session_state.fl_info}


背景：

{st.session_state.bg_info}



【剧情档案】

{st.session_state.bible_info}



【文风档案】

{st.session_state.style_info}



结尾追加：

【档案更新】

记录：

本段事件

人物关系变化

新增设定

伏笔

"""



from google.generativeai.types import HarmCategory, HarmBlockThreshold

# 修正后的安全配置定义
safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# 重新初始化模型
model = genai.GenerativeModel(
    "gemini-2.5-pro", # 建议使用稳定版本，如 1.5-pro
    system_instruction=system_prompt,
    safety_settings=safety_settings
)







# =====================
# 历史显示
# =====================


for msg in st.session_state.chat_history[-10:]:

    with st.chat_message(
        "user"
        if msg["role"]=="user"
        else "assistant"
    ):

        st.markdown(
            msg["text"]
        )






# =====================
# 创作
# =====================


prompt=st.text_area(
    "✨ 输入剧情指令:",
    key="input",
    height=200
)



c1,c2=st.columns(2)



if c1.button(
    "✨立即创作"
):


    if prompt:


        history=[

            {
                "role":
                m["role"],

                "parts":
                [
                    m["text"]
                ]

            }

            for m in st.session_state.chat_history[-5:]

        ]


        chat=model.start_chat(
            history=history
        )


        with st.spinner(
            "正在创作..."
        ):


            response=safe_send(
                chat,
                prompt
            )


        if response:


            text=response.text


            st.session_state.chat_history.append(
                {
                    "role":"user",
                    "text":prompt
                }
            )


            st.session_state.chat_history.append(
                {
                    "role":"model",
                    "text":text
                }
            )



            if "【档案更新】" in text:


                update=text.split(
                    "【档案更新】",
                    1
                )[1]


                st.session_state.bible_info += (
                    "\n"+update
                )



            sync_data()

            st.rerun()





if c2.button(
    "↩️ 撤回上一段"
):

    if len(
        st.session_state.chat_history
    )>=2:


        st.session_state.chat_history=(
            st.session_state.chat_history[:-2]
        )


        sync_data()

        st.rerun()
